import os
import pytest


def test_reference_document_lifecycle_full_integration(tmp_path):
    """Full-stack integration test exercising upload -> ingestion -> persistence -> retrieval.

    This test requires the full stack: a running database reachable via
    `DATABASE_URL`, a running Celery worker and broker, and a writable
    `REFERENCE_STORAGE_PATH`. It uploads a DOCX from `backend/data/references`
    and waits for the worker to ingest and persist chunks.
    """

    # locate a real reference DOCX in repo
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    data_dir = os.path.join(repo_root, "data", "references")
    if not os.path.isdir(data_dir):
        pytest.skip("No backend/data/references directory with DOCX files; skip integration test")

    docx_files = [f for f in os.listdir(data_dir) if f.lower().endswith(".docx")] 
    # If no docx exists, create a small sample reference doc for testing
    if not docx_files:
        try:
            from docx import Document

            sample_path = os.path.join(data_dir, "sample_reference.docx")
            doc = Document()
            doc.add_heading("Cardiovascular Agents", level=1)
            doc.add_heading("Velantine", level=2)
            doc.add_paragraph("Velantine is used for ...")
            doc.add_paragraph("Dose: 50 mg once daily")
            doc.add_paragraph("Contraindications: ...")
            doc.add_heading("Cordizem-XR", level=2)
            doc.add_paragraph("Cordizem-XR is useful for ...")
            doc.add_paragraph("Dose: 100 mg once daily")
            doc.add_heading("Mirosartan", level=2)
            doc.add_paragraph("Mirosartan notes ...")
            os.makedirs(data_dir, exist_ok=True)
            doc.save(sample_path)
            docx_files = [os.path.basename(sample_path)]
        except Exception:
            pytest.skip("No .docx files in backend/data/references and cannot create sample docx; skip integration test")

    src_file = os.path.join(data_dir, docx_files[0])

    # This is a full-stack test: use a temp folder for saved uploads
    # The app should save incoming files to `REFERENCE_STORAGE_PATH`.
    upload_dir = tmp_path / "uploads"
    upload_dir.mkdir()
    os.environ["REFERENCE_STORAGE_PATH"] = str(upload_dir)

    # Start app client and upload file via API
    from backend.app.main import app
    from fastapi.testclient import TestClient

    client = TestClient(app)

    with open(src_file, "rb") as f:
        files = {"file": (os.path.basename(src_file), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = client.post("/api/v1/references", files=files)

    assert resp.status_code == 202, f"upload failed: {resp.text}"
    body = resp.json()
    ref_id = body.get("reference_id")
    assert ref_id is not None

    # Poll DB for the document to reach READY (a real Celery worker should process the task)
    from backend.app.db.session import SessionLocal
    from backend.app.db.models.reference_document import ReferenceDocument, ReferenceStatus
    import time

    db = SessionLocal()
    try:
        timeout = int(os.getenv("FULL_STACK_TIMEOUT", "60"))
        start = time.time()
        doc = None
        while time.time() - start < timeout:
            doc = db.query(ReferenceDocument).filter(ReferenceDocument.id == ref_id).first()
            if doc is not None and doc.status == ReferenceStatus.READY:
                break
            time.sleep(1)
        assert doc is not None, "ReferenceDocument record not found (full-stack mode)"
        assert doc.status == ReferenceStatus.READY, f"Expected READY but was {getattr(doc, 'status', None)}"
    finally:
        db.close()

    # Verify DB persisted document and chunks
    from backend.app.db.session import SessionLocal
    from backend.app.db.repositories.reference_repository import ReferenceRepository
    from backend.app.db.models.reference_document import ReferenceDocument, ReferenceStatus

    db = SessionLocal()
    try:
        ref = db.query(ReferenceDocument).filter(ReferenceDocument.id == ref_id).first()
        assert ref is not None, "ReferenceDocument record not found"
        assert ref.status == ReferenceStatus.READY, f"Expected READY but was {ref.status}"

        repo = ReferenceRepository(db)
        chunks = repo.get_chunks_by_reference(ref_id)
        assert chunks and len(chunks) > 0, "No chunks persisted for reference"

        # Validate chunk properties and embeddings
        for ch in chunks:
            assert ch.content and ch.content.strip(), "Empty chunk content"
            emb = getattr(ch, "embedding", None)
            assert emb is not None, "Missing embedding on persisted chunk"
            # embedding stored as list in SQLite path
            assert isinstance(emb, (list, tuple)), "Embedding stored in unexpected format"
            assert len(emb) == 384, f"Unexpected embedding dimension: {len(emb)}"
            # metadata/source
            meta = getattr(ch, "meta", None)
            assert meta is not None and meta.get("source") == os.path.basename(src_file)

        # Retrieval: pick a phrase from first chunk and retrieve evidence
        from backend.app.domain.schemas import MedicationEntity
        from backend.app.domain.evidence_service import EvidenceService

        first_chunk_text = chunks[0].content
        # pick a phrase of several words found in the chunk
        tokens = [t for t in first_chunk_text.split() if len(t) > 3]
        assert tokens, "No suitable token in chunk to build query"
        query_phrase = " ".join(tokens[:4])

        entity = MedicationEntity(medication_name=query_phrase, dose_value=None)
        evsvc = EvidenceService()
        evidence = evsvc.retrieve_evidence(entity, ref_id, top_k=5, similarity_threshold=0.0)

        assert evidence is not None
        assert len(evidence) > 0, "Retrieval returned no evidence"

        # Verify retrieved evidence documents point to the same reference
        for ev in evidence:
            assert ev.document_id == ref_id
            assert ev.text and query_phrase.split()[0] in ev.text

    finally:
        db.close()
