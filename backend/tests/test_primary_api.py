from io import BytesIO
import os
from types import SimpleNamespace
from docx import Document


def make_docx_bytes():
    doc = Document()
    doc.add_paragraph("Primary Title")
    doc.add_paragraph("Patient instructed to take Velantine 50 mg once daily.")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def test_upload_primary_creates_job(client, monkeypatch, tmp_path):
    # The endpoint selects the first primary document from the configured directory.
    storage = tmp_path / "primary"
    storage.mkdir()
    (storage / "primary.docx").write_bytes(make_docx_bytes().getvalue())
    os.environ["PRIMARY_STORAGE_PATH"] = str(storage)

    # monkeypatch DocumentRepository used in endpoint
    import backend.app.api.v1.endpoints.primary_documents as pdmod

    class FakeDocRepo:
        def __init__(self, db):
            pass

        def create_document(self, filename, user_id, mime_type, content_hash):
            from types import SimpleNamespace
            return SimpleNamespace(id=1, filename=filename, user_id=user_id, mime_type=mime_type, content_hash=content_hash, pages=None, created_at=None)

    monkeypatch.setattr(pdmod, "DocumentRepository", FakeDocRepo)

    # monkeypatch JobRepository used by endpoint
    import backend.app.api.v1.endpoints.primary_documents as pdmod2

    class FakeJobRepo:
        def __init__(self, db):
            pass

        def create_job(self, user_id, primary_document_id, reference_document_id):
            from types import SimpleNamespace
            return SimpleNamespace(id=1, primary_document_id=primary_document_id, reference_document_id=reference_document_id, user_id=user_id, status="QUEUED", current_stage=None, created_at=None)

    monkeypatch.setattr(pdmod2, "JobRepository", FakeJobRepo)

    # fake Celery task
    import backend.app.api.v1.endpoints.primary_documents as pdmod3

    def fake_delay(*args, **kwargs):
        from types import SimpleNamespace
        return SimpleNamespace(id="dummy-task-id")

    monkeypatch.setattr(pdmod3, "primary_document_pipeline", SimpleNamespace(delay=fake_delay))

    resp = client.post("/api/v1/primary-documents?reference_id=1")
    assert resp.status_code == 202
    data = resp.json()
    assert data["document_id"] == 1
    assert data["status"] == "queued"
