from io import BytesIO
from docx import Document


def make_docx_bytes():
    doc = Document()
    doc.add_paragraph("Test Title")
    doc.add_paragraph("First paragraph of the document.")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def test_upload_docx_creates_reference(client):
    """Upload a valid DOCX file and expect a 202 response with a created reference id and ingest job id."""
    bio = make_docx_bytes()
    files = {"file": ("test.docx", bio, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    resp = client.post("/api/v1/references", files=files)
    assert resp.status_code == 202
    data = resp.json()
    assert data["reference_id"] == 1
    assert data["ingest_job_id"] == "dummy-task-id"


def test_upload_invalid_extension_rejected(client):
    """Upload a file with an unsupported extension and expect a 400 Bad Request response."""
    bad = BytesIO(b"not a docx")
    files = {"file": ("bad.txt", bad, "text/plain")}
    resp = client.post("/api/v1/references", files=files)
    assert resp.status_code == 400
